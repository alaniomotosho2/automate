import docx
doc = docx.Document('tay.docx')


########### it read and output column by column
table = doc.tables[0]
print("Column 1:")
for i in range(len(table.rows)):
	print(table.rows[i].cells[0].paragraphs[0].text)
print("Column 2:")
for i in range(len(table.rows)):
	print(table.rows[i].cells[1].paragraphs[0].text)
print("Column 3:")
for i in range(len(table.rows)):
	print(table.rows[i].cells[2].paragraphs[0].text)