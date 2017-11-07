from docx import Document
from docx.shared import Inches

document = Document()

#document = Document('testImage.docx')





document.add_picture('download.jpeg',width=Inches(1.25))
document.save('testImage.docx')