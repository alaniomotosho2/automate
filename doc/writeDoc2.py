from docx import Document

document = Document()

document = Document('testImage.docx')
p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold words').bold = True
p.add_run(' and italics.').italic = True
document.save('testImage.docx')


document = Document('testImage.docx')
document.add_heading('Lets talk about Python language', level=1)
document.add_paragraph('First lets see the Python logo', style='ListBullet')
document.save('testImage.docx')